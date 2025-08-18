"""
文件处理模块
支持Word、PDF、Excel等文件的解析和分析
"""

import io
import os
import docx
import PyPDF2
import pandas as pd
from typing import Dict, List, Optional, Any
import streamlit as st


class FileProcessor:
    """文件处理器"""

    def __init__(self):
        self.supported_types = {
            'docx': 'Word文档',
            'pdf': 'PDF文档',
            'xlsx': 'Excel表格',
            'xls': 'Excel表格',
            'txt': '文本文件'
        }

    def get_file_type(self, file) -> str:
        """获取文件类型"""
        if hasattr(file, 'name'):
            filename = file.name
        else:
            filename = str(file)

        return filename.split('.')[-1].lower()

    def is_supported(self, file) -> bool:
        """检查是否支持该文件类型"""
        file_type = self.get_file_type(file)
        return file_type in self.supported_types

    def process_file(self, file) -> Dict[str, Any]:
        """
        处理上传的文件

        Args:
            file: streamlit上传的文件对象

        Returns:
            包含文件信息和内容的字典
        """
        try:
            file_type = self.get_file_type(file)

            if not self.is_supported(file):
                return {
                    'success': False,
                    'error': f'不支持的文件类型: {file_type}',
                    'content': None
                }

            # 根据文件类型处理
            if file_type == 'docx':
                content = self._process_word(file)
            elif file_type == 'pdf':
                content = self._process_pdf(file)
            elif file_type in ['xlsx', 'xls']:
                content = self._process_excel(file)
            elif file_type == 'txt':
                content = self._process_text(file)
            else:
                return {
                    'success': False,
                    'error': f'暂不支持处理 {file_type} 文件',
                    'content': None
                }

            return {
                'success': True,
                'file_name': file.name,
                'file_type': file_type,
                'file_size': len(file.getvalue()),
                'content': content,
                'summary': self._generate_summary(content, file_type)
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'文件处理失败: {str(e)}',
                'content': None
            }

    def _process_word(self, file) -> str:
        """处理Word文档"""
        try:
            doc = docx.Document(file)
            content = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())

            # 提取表格内容
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_content.append(cell.text.strip())
                    if row_content:
                        table_content.append(' | '.join(row_content))

                if table_content:
                    content.append('\n表格内容:\n' + '\n'.join(table_content))

            return '\n\n'.join(content)

        except Exception as e:
            raise Exception(f"Word文档处理失败: {str(e)}")

    def _process_pdf(self, file) -> str:
        """处理PDF文档"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            content = []

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"第{page_num + 1}页:\n{text.strip()}")
                except Exception as e:
                    content.append(f"第{page_num + 1}页: 无法提取文本 ({str(e)})")

            return '\n\n'.join(content)

        except Exception as e:
            raise Exception(f"PDF文档处理失败: {str(e)}")

    def _process_excel(self, file) -> str:
        """处理Excel文档"""
        try:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file)
            content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file, sheet_name=sheet_name)

                # 生成工作表摘要
                sheet_info = [
                    f"工作表: {sheet_name}",
                    f"行数: {len(df)}",
                    f"列数: {len(df.columns)}",
                    f"列名: {', '.join(df.columns.astype(str))}"
                ]

                # 添加前几行数据作为示例
                if len(df) > 0:
                    sheet_info.append("\n前5行数据:")
                    sheet_info.append(df.head().to_string())

                # 添加数据统计
                if len(df) > 0:
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        sheet_info.append("\n数值列统计:")
                        sheet_info.append(df[numeric_cols].describe().to_string())

                content.append('\n'.join(sheet_info))

            return '\n\n' + '=' * 50 + '\n\n'.join(content)

        except Exception as e:
            raise Exception(f"Excel文档处理失败: {str(e)}")

    def _process_text(self, file) -> str:
        """处理文本文件"""
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            content = None

            for encoding in encodings:
                try:
                    file.seek(0)  # 重置文件指针
                    content = file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise Exception("无法识别文件编码")

            return content

        except Exception as e:
            raise Exception(f"文本文件处理失败: {str(e)}")

    def _generate_summary(self, content: str, file_type: str) -> str:
        """生成文件内容摘要"""
        if not content:
            return "文件内容为空"

        content_length = len(content)
        word_count = len(content.split())
        line_count = len(content.split('\n'))

        summary = [
            f"文件类型: {self.supported_types.get(file_type, file_type)}",
            f"内容长度: {content_length} 字符",
            f"单词数量: {word_count}",
            f"行数: {line_count}"
        ]

        # 添加内容预览
        preview = content[:200] + "..." if len(content) > 200 else content
        summary.append(f"内容预览: {preview}")

        return '\n'.join(summary)

    def extract_key_info(self, content: str, file_type: str) -> Dict[str, Any]:
        """
        提取文件关键信息用于AI分析

        Args:
            content: 文件内容
            file_type: 文件类型

        Returns:
            关键信息字典
        """
        key_info = {
            'content_summary': self._generate_summary(content, file_type),
            'key_points': [],
            'questions': [],
            'data_insights': []
        }

        # 提取关键点（简单的关键词识别）
        keywords = ['学习', '课程', '成绩', '考试', '作业', '项目', '计划', '目标', '问题', '困难', '压力', '焦虑']

        for keyword in keywords:
            if keyword in content:
                # 找到包含关键词的句子
                sentences = content.split('。')
                relevant_sentences = [s.strip() for s in sentences if keyword in s and len(s.strip()) > 5]
                if relevant_sentences:
                    key_info['key_points'].extend(relevant_sentences[:2])  # 最多取2个相关句子

        # 识别可能的问题或疑问
        question_indicators = ['？', '?', '怎么', '如何', '为什么', '是否', '能否']
        sentences = content.split('。')

        for sentence in sentences:
            for indicator in question_indicators:
                if indicator in sentence and len(sentence.strip()) > 5:
                    key_info['questions'].append(sentence.strip())
                    break

        # 去重并限制数量
        key_info['key_points'] = list(set(key_info['key_points']))[:5]
        key_info['questions'] = list(set(key_info['questions']))[:3]

        return key_info


def create_file_upload_section():
    """创建文件上传区域"""
    st.markdown("### 📎 文件上传分析")

    processor = FileProcessor()

    # 显示支持的文件类型
    with st.expander("📋 支持的文件类型", expanded=False):
        for ext, desc in processor.supported_types.items():
            st.write(f"• **{ext.upper()}**: {desc}")

    # 文件上传
    uploaded_file = st.file_uploader(
        "上传文件进行分析",
        type=list(processor.supported_types.keys()),
        help="上传相关文档，AI将结合文档内容为您提供更准确的建议"
    )

    if uploaded_file is not None:
        # 处理文件
        with st.spinner("🔍 正在分析文件内容..."):
            result = processor.process_file(uploaded_file)

        if result['success']:
            # 显示文件信息
            st.success(f"✅ 文件 '{result['file_name']}' 上传成功!")

            with st.expander("📊 文件信息", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**文件名:** {result['file_name']}")
                    st.write(f"**文件类型:** {result['file_type'].upper()}")
                with col2:
                    st.write(f"**文件大小:** {result['file_size']} 字节")
                    st.write(f"**处理状态:** 成功")

                st.markdown("**文件摘要:**")
                st.text(result['summary'])

            # 显示关键信息
            key_info = processor.extract_key_info(result['content'], result['file_type'])

            if key_info['key_points'] or key_info['questions']:
                with st.expander("🔑 关键信息提取", expanded=True):
                    if key_info['key_points']:
                        st.markdown("**📌 关键点:**")
                        for i, point in enumerate(key_info['key_points'], 1):
                            st.write(f"{i}. {point}")

                    if key_info['questions']:
                        st.markdown("**❓ 发现的问题:**")
                        for i, question in enumerate(key_info['questions'], 1):
                            st.write(f"{i}. {question}")

            # 将文件内容存储到session state中供AI使用
            st.session_state.uploaded_file_content = {
                'file_name': result['file_name'],
                'content': result['content'],
                'key_info': key_info,
                'summary': result['summary']
            }

            st.info("💡 文件已上传成功！AI助手现在可以基于您上传的文件内容提供更精准的建议。")

        else:
            st.error(f"❌ 文件处理失败: {result['error']}")

    return uploaded_file